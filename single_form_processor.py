from docx import Document
from docx.shared import RGBColor
import json

from google import genai
from google.genai import types
import pathlib
import os
import PIL


# Set up the API key and client
api_key = "AIzaSyAQtSELrheQ1szjk8oBBJAxq0RXgRdpWvg"
client = genai.Client(api_key=api_key)

# Replace this with your actual local PDF path
filepath = pathlib.Path("/Users/thevedantsingh/Desktop/SLS/D004 - UD 105 - Specific Denials LA - Kyle Answer.pdf")

# Read the local PDF and send it to Gemini
prompt = 'Give me all information you have about the case'
response = client.models.generate_content(
    model="gemini-2.0-flash-exp",
    contents=[
        types.Part.from_bytes(
            data=filepath.read_bytes(),
            mime_type='application/pdf',
        ),
        prompt
    ]
)

# Print the result
cont = response.text

def analyze_template(template_path):
    
    doc = Document(template_path)
    template_data = []
    
    for para in doc.paragraphs:
        para_info = {
            "style": para.style.name,
            "runs": [],
            "placeholders": []
        }
        
        for run in para.runs:
            # Capture formatting
            fmt = {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "color": run.font.color.rgb if run.font.color.rgb else None
            }
            para_info["runs"].append(fmt)
            
            # Identify placeholders
            if '#' in run.text and run.font.color.rgb == RGBColor(0xFF,0,0):
                para_info["placeholders"].append({
                    "placeholder": run.text.strip(),
                    "position": len(para_info["runs"])-1
                })
                
        template_data.append(para_info)
    
    with open("template_blueprint.json", "w") as f:
        json.dump(template_data, f, indent=2)


analyze_template("/Users/thevedantsingh/Desktop/SLS/Jury Instruction Template MASTER - 2022 - LAFLA.docx")
import json
with open("template_blueprint.json") as f:
    data = json.load(f)

with open("output.txt", "w", encoding="utf-8") as out:
    for para in data:
        out.write(json.dumps(para, ensure_ascii=False))  # Converts dict to string
        out.write("\n")  # Newline after each paragraph

def generate_content():
    api_key = "AIzaSyAQtSELrheQ1szjk8oBBJAxq0RXgRdpWvg"
    client = genai.Client(api_key=api_key)
    myfile = client.files.upload(file="/Users/thevedantsingh/Desktop/SLS/output.txt")
    # Build prompt with formatting context
    prompt = f"""Legal Document Generation Task:
    Fill in the CACI template for this given CACI form using the provided context: "CACI Number": "101", "Jury Instructions": "Overview of Trial". Use this as context: {cont} to fill the form placeholders. Use the context and fill in placeholder such as defendant/plaintiff/city/address. Make your best guess for the placeholders. I want to see all placeholders filled. Dont output anything else. Keep the output in the same format as the input. JSON format, with the same attributes as the provided doc, you only need to change the text as you see fit.
    Use the same attributes as the provided text with the same styling. I need the same format as input. JSON format with the style, run, text, everything in style formatting. Fill the entire CACI form.  I need the entire form bro, i dont get what you are not understanding. I only need the CACI 101 form: Overview of a trial. I dont need superflous stuff
    Ignore any lines of this form. I only need the CACI 101. I need everything in one json file, i dont need separate JSON records.
    "style": "normal",
    "runs": [
      
        "text": "",
        "bold": null,
        "italic": null,
        "color": null
      
    ],
    "placeholders": []
  ,
    Output:"""
    
    response = client.models.generate_content(
            model="gemini-2.0-flash-exp",
            contents=[
                prompt, myfile
            ]
    )
    
    return response.text


import json
from docx import Document
from docx.shared import RGBColor
import shutil
import re

def extract_json_objects(text):
    print(text)
    match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
    if not match:
        raise ValueError("No JSON block found.")

    json_block = match.group(1)
    json_block = re.sub(r'\}\s*,?\s*\{', '},{', json_block)
    json_block = f"[{json_block}]"

    try:
        parsed = json.loads(json_block)
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON: {e}")

    return parsed

def get_first_10_records(text):
    data = extract_json_objects(text)
    return data

def color_from_hex(hexstr):
    if isinstance(hexstr, list):
        return RGBColor(*hexstr)
    if not hexstr or hexstr.lower() == "auto":
        return None
    hexstr = hexstr.lstrip('#')
    return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))

def json_to_docx(json_data, docx_path, template_path=None):
    if template_path:
        shutil.copyfile(template_path, docx_path)
        doc = Document(docx_path)
        for element in reversed(doc.paragraphs):
            element.clear()
    else:
        doc = Document()

    for para in json_data:
        new_para = doc.add_paragraph(style=para.get("style", "Normal"))
        for run_info in para["runs"]:
            run = new_para.add_run(run_info["text"])
            run.bold = run_info.get("bold", False)
            run.italic = run_info.get("italic", False)
            if "color" in run_info:
                run.font.color.rgb = color_from_hex(run_info["color"])
    print("YESYSYYSYS")

    doc.save(docx_path)
try:
    result = generate_content()
    with open("debug.txt", "w") as file:
        file.write(str(result))
    data = get_first_10_records(result)
    print(data)
    print("TYPE = ", type(data))

    for item in data:
        if "style" in item and isinstance(item["style"], str):
            item["style"] = item["style"].capitalize()
    print(data)    
    json_to_docx(
        data,
        "final_output.docx",
    )
    print('DONE')

except json.JSONDecodeError as e:
    print(f"Invalid JSON: {str(e)}")
    print("Raw LLM output:")
    print(result)
except Exception as e:
    print(f"Error: {str(e)}")