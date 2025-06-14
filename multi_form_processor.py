from docx import Document
from docx.shared import RGBColor, Pt
import json
import re
import shutil
import pathlib

from google import genai
from google.genai import types

# Set up the API key and client
api_key = "ADD API KEY"
client = genai.Client(api_key=api_key)

def analyze_pdf_with_gemini(pdf_path, prompt_text, api_key):
    """Send PDF to Gemini and get analysis results"""
    client = genai.Client(api_key=api_key)
    filepath = pathlib.Path(pdf_path)
    
    response = client.models.generate_content(
        model="gemini-2.0-flash-exp",
        contents=[
            types.Part.from_bytes(
                data=filepath.read_bytes(),
                mime_type='application/pdf',
            ),
            prompt_text
        ]
    )
    
    return response.text

def analyze_template(template_path):
    """Analyze Word template to identify placeholders and styles"""
    doc = Document(template_path)
    template_data = []
    
    for para in doc.paragraphs:
        para_info = {
            "style": para.style.name,
            "runs": [],
            "placeholders": []
        }
        
        for run in para.runs:
            # Capture formatting
            fmt = {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "color": run.font.color.rgb if run.font.color and run.font.color.rgb else None
            }
            para_info["runs"].append(fmt)
            
            # Identify placeholders
            if '#' in run.text and run.font.color and run.font.color.rgb == RGBColor(0xFF, 0, 0):
                para_info["placeholders"].append({
                    "placeholder": run.text.strip(),
                    "position": len(para_info["runs"])-1
                })
                
        template_data.append(para_info)
    
    with open("template_blueprint.json", "w") as f:
        json.dump(template_data, f, indent=2)
    
    return template_data

def extract_json_objects(text):
    """Extract JSON objects from text that may be wrapped in code blocks"""
    match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
    if match:
        json_block = match.group(1)
    else:
        match = re.search(r'```\s*(.*?)\s*```', text, re.DOTALL)
        if match:
            json_block = match.group(1)
        else:
            json_block = text.strip()
    
    json_block = re.sub(r'\}\s*,?\s*\{', '},{', json_block)
    if not json_block.startswith('['):
        json_block = f"[{json_block}]"
    if not json_block.endswith(']'):
        json_block = f"{json_block}]"
    
    try:
        parsed = json.loads(json_block)
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON: {e}")

    return parsed

def color_from_hex(hexstr):
    """Convert various color formats to RGBColor"""
    if isinstance(hexstr, list):
        if len(hexstr) == 3:
            return RGBColor(*hexstr)
        return None
    if not hexstr or hexstr.lower() == "auto":
        return None
    if isinstance(hexstr, str):
        hexstr = hexstr.lstrip('#')
        if len(hexstr) == 6:
            return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))
    return None

def generate_content(caci_forms):
    """Generate content for multiple CACI forms and combine them into one document"""
    api_key = "ADD API KEY"
    client = genai.Client(api_key=api_key)
    
    # Get the case information from the PDF
    filepath = pathlib.Path("UD 105 Path")
    prompt = 'Give me all information you have about the case'
    response = client.models.generate_content(
        model="gemini-2.0-flash-exp",
        contents=[
            types.Part.from_bytes(
                data=filepath.read_bytes(),
                mime_type='application/pdf',
            ),
            prompt
        ]
    )
    cont = response.text
    
    # Upload the template structure
    myfile = client.files.upload(file="Template txt file; output.txt")
    
    # Create a new document to hold all forms
    combined_doc = Document()
    
    # Process each form identified from the LA CIV 244
    for i, form in enumerate(caci_forms):
        caci_number = form["CACI Number"]
        jury_instructions = form["Jury Instructions"]
        
        print(f"Processing {i+1}/{len(caci_forms)}: CACI {caci_number} - {jury_instructions}")
        
        # Build prompt with formatting context for each form
        prompt = f"""Legal Document Generation Task:
        Fill in the CACI template for this given CACI form using the provided context: "CACI Number": "{caci_number}", "Jury Instructions": "{jury_instructions}". Use this as context: {cont} to fill the form placeholders. Use the context and fill in placeholder such as defendant/plaintiff/city/address. Make your best guess for the placeholders. I want to see all placeholders filled. Dont output anything else. Keep the output in the same format as the input. JSON format, with the same attributes as the provided doc, you only need to change the text as you see fit.
        Use the same attributes as the provided text with the same styling. I need the same format as input. JSON format with the style, run, text, everything in style formatting. Fill the entire CACI form. I need the entire form, with all placeholders filled in.
        Ignore any lines of this form 
        "style": "normal",
        "runs": [
          
            "text": "",
            "bold": null,
            "italic": null,
            "color": null
          
        ],
        "placeholders": []
      ,
        Output:"""
        
        response = client.models.generate_content(
                model="gemini-2.0-flash-exp",
                contents=[
                    prompt, myfile
                ]
        )
        
        result_text = response.text
        
        # Extract JSON data
        try:
            form_data = extract_json_objects(result_text)
            
            # Fix style capitalization if needed
            for item in form_data:
                if "style" in item and isinstance(item["style"], str):
                    item["style"] = item["style"].capitalize()
            
            # Add form title to the combined document
            title_para = combined_doc.add_paragraph()
            title_run = title_para.add_run(f"CACI {caci_number}: {jury_instructions}")
            title_run.bold = True
            title_run.font.size = Pt(14)
            
            # Add form content
            for para in form_data:
                new_para = combined_doc.add_paragraph(style=para.get("style", "Normal"))
                for run_info in para["runs"]:
                    run = new_para.add_run(run_info["text"])
                    run.bold = run_info.get("bold", False)
                    run.italic = run_info.get("italic", False)
                    if "color" in run_info and run_info["color"]:
                        run.font.color.rgb = color_from_hex(run_info["color"])
            
            # Add page break after each form (except the last one)
            if i < len(caci_forms) - 1:
                combined_doc.add_page_break()
                
            print(f'✓ Added CACI {caci_number}')
            
        except Exception as e:
            print(f"✗ Error processing CACI {caci_number}: {str(e)}")
            # Add error note to document
            error_para = combined_doc.add_paragraph()
            error_run = error_para.add_run(f"Error processing CACI {caci_number}: {str(e)}")
            error_run.font.color.rgb = RGBColor(255, 0, 0)
            
            # Add page break after error (except if it's the last form)
            if i < len(caci_forms) - 1:
                combined_doc.add_page_break()
    
    # Save the combined document
    output_path = "final_output.docx"
    combined_doc.save(output_path)
    print(f"\nAll forms combined into {output_path}")
    
    return output_path

def main():
    # Set up templates and directories
    template_path = "Master LAFLA Template"
    template_data = analyze_template(template_path)
    
    # Save template data to a file for reference by the Gemini model
    with open("output.txt", "w", encoding="utf-8") as out:
        for para in template_data:
            out.write(json.dumps(para, ensure_ascii=False))  # Converts dict to string
            out.write("\n")  # Newline after each paragraph
    
    # Identify forms that need to be filled
    api_key = "ADD API KEY"
    pd_path = "LA CIV PDF"
    prompt = f"""Analyse the given LA CIV 244 submission and then identify all the forms that need to be filled based on whatever defenses/forms
    are ticked/checked in the LA CIV. Then return your output in json format, with each record having two keys, the keys being CACI Number and Jury Instructions. Make sure every form that is checked/ticked is included. I need no other output."""
    
    result = analyze_pdf_with_gemini(pd_path, prompt, api_key)
    
    # Extract JSON from result
    if "```json" in result:
        # Extract content between ```json and the last ```
        json_text = result.split("```json")[1].split("```")[0].strip()
    elif "```" in result:
        # Extract content between first ``` and the last ```
        json_text = result.split("```")[1].split("```")[0].strip()
    else:
        json_text = result.strip()
    
    # Parse JSON to get list of forms to process
    try:
        caci_forms = json.loads(json_text)
        print(f"Found {len(caci_forms)} CACI forms to process:")
        for form in caci_forms:
            print(f"- CACI {form['CACI Number']}: {form['Jury Instructions']}")
        
        # Process all forms and create a single combined document
        output_file = generate_content(caci_forms)
        print(f"\nProcessing complete. All forms have been combined into {output_file}")
    except json.JSONDecodeError as e:
        print(f"Invalid JSON: {str(e)}")
        print("Raw output:")
        print(result)
    except Exception as e:
        print(f"Error during processing: {str(e)}")

if __name__ == "__main__":
    main()