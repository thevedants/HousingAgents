import os
import re
import json
import tempfile
import shutil
import logging
import time
from flask import Flask, request, render_template_string, send_file
from docx import Document
from docx.shared import RGBColor, Pt
import pathlib
import datetime

# Import Google GenAI
from google import genai
from google.genai import types

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# Create a permanent directory for response logs
RESPONSE_LOGS_DIR = "gemini_responses"
os.makedirs(RESPONSE_LOGS_DIR, exist_ok=True)

# HTML template for the form
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CACI Form Generator</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <style>
        body {
            background-color: #f5f5f5;
            font-family: Arial, sans-serif;
        }
        .container {
            max-width: 800px;
            margin: 50px auto;
            background-color: white;
            border-radius: 10px;
            padding: 30px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }
        .header {
            text-align: center;
            margin-bottom: 30px;
        }
        .upload-section {
            border: 2px dashed #ccc;
            border-radius: 8px;
            padding: 20px;
            text-align: center;
            margin-bottom: 20px;
        }
        .submit-btn {
            width: 100%;
            padding: 12px;
            font-size: 16px;
            margin-top: 20px;
        }
        pre {
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
            font-size: 12px;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>CACI Form Generator</h1>
            <p class="text-muted">Upload your legal documents to generate filled CACI forms</p>
        </div>

        <form action="/process" method="post" enctype="multipart/form-data">
            <div class="row">
                <div class="col-md-6">
                    <div class="upload-section">
                        <h4>LA CIV 244 PDF</h4>
                        <p>Upload the LA CIV 244 document</p>
                        <input type="file" class="form-control" name="laciv_file" accept=".pdf" required>
                    </div>
                </div>
                <div class="col-md-6">
                    <div class="upload-section">
                        <h4>Case Details PDF</h4>
                        <p>Upload the UD 105 document</p>
                        <input type="file" class="form-control" name="case_file" accept=".pdf" required>
                    </div>
                </div>
            </div>

            <div class="upload-section mt-3">
                <h4>CACI Template (DOCX)</h4>
                <p>Upload the Jury Instruction Template document</p>
                <input type="file" class="form-control" name="template_file" accept=".txt" required>
            </div>

            <div class="form-check mt-3">
                <input class="form-check-input" type="checkbox" id="save-responses" name="save_responses" checked>
                <label class="form-check-label" for="save-responses">
                    Save Gemini API responses to files
                </label>
            </div>
            
            <div class="form-check mt-2">
                <input class="form-check-input" type="checkbox" id="debug-mode" name="debug_mode" checked>
                <label class="form-check-label" for="debug-mode">
                    Show processing details
                </label>
            </div>

            <button type="submit" class="btn btn-primary submit-btn">Generate CACI Forms</button>
        </form>
    </div>
</body>
</html>
"""

@app.route('/', methods=['GET'])
def index():
    # Render the HTML form
    return render_template_string(HTML_TEMPLATE)

@app.route('/process', methods=['POST'])
def process():
    # Create a temporary directory for all processing
    temp_dir = tempfile.mkdtemp()
    debug_mode = 'debug_mode' in request.form
    save_responses = 'save_responses' in request.form
    debug_output = []
    
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    response_log_files = []

    
    def save_response_to_file(response_text, label):
        import json
        if save_responses:
            filename = f"{timestamp}_{label}.txt"
            filepath = os.path.join(RESPONSE_LOGS_DIR, filename)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(response_text)
            response_log_files.append(filepath)
            debug_output.append(f"Saved {label} response to: {filepath}")
            return filepath
        return None
    
    try:
        # Save uploaded files to temp directory
        
        laciv_file = request.files['laciv_file']
        case_file = request.files['case_file']
        template_file = request.files['template_file']
        
        laciv_path = os.path.join(temp_dir, secure_filename(laciv_file.filename))
        case_path = os.path.join(temp_dir, secure_filename(case_file.filename))
        template_path = os.path.join(temp_dir, secure_filename(template_file.filename))
        
        laciv_file.save(laciv_path)
        case_file.save(case_path)
        template_file.save(template_path)
        
        debug_output.append(f"Files saved successfully")
        
        # Output paths
        output_path = os.path.join(temp_dir, 'final_output.docx')
        
        # Set up the API key
        api_key = "AIzaSyAQtSELrheQ1szjk8oBBJAxq0RXgRdpWvg"  # Replace with your actual API key
        client = genai.Client(api_key=api_key)
        filepath = pathlib.Path(case_path)
        prompt = 'Give me all information you have about the case'
        
        debug_output.append("Getting case information from PDF...")
        
        try:
            response = client.models.generate_content(
                model="gemini-2.0-flash-exp",
                contents=[
                    types.Part.from_bytes(
                        data=filepath.read_bytes(),
                        mime_type='application/pdf',
                    ),
                    prompt
                ]
            )
            case_info = response.text
            debug_output.append(f"Case information received ({len(case_info)} characters)")
            
            # Save the case information response
            save_response_to_file(case_info, "case_info")
            
        except Exception as e:
            error_msg = f"Error getting case info: {str(e)}"
            debug_output.append(error_msg)
            logger.error(error_msg)
            raise
        
        # Now try to process template in simpler way
        # APPROACH 1: Try generating JSON with explicit structure
        debug_output.append("Attempting to generate form with JSON structure...")
        
        # First try with a more explicit JSON structure
        json_prompt = f"""Legal Document Generation Task:
    Fill in the CACI template for this given CACI form using the provided context: "CACI Number": "101", "Jury Instructions": "Overview of Trial". Use this as context: {case_info} to fill the form placeholders. Use the context and fill in placeholder such as defendant/plaintiff/city/address. Make your best guess for the placeholders. I want to see all placeholders filled. Dont output anything else. Keep the output in the same format as the input. JSON format, with the same attributes as the provided doc, you only need to change the text as you see fit.
    Use the same attributes as the provided text with the same styling. I need the same format as input. JSON format with the style, run, text, everything in style formatting. Fill the entire CACI form.  I need the entire form bro, i dont get what you are not understanding. I only need the CACI 101 form: Overview of a trial. I dont need superflous stuff
    Ignore any lines of this form. I only need the CACI 101:Overview of Trial. Your response should be one JSON file only with records being separated by commas. Only the form style json, no placeholders, nothing apart from the filled form w style details
    "style": "normal",
    "runs": [
      
        "text": "",
        "bold": null,
        "italic": null,
        "color": null
      
    ]
    Output:"""
        
        try:
            myfile = client.files.upload(file=template_path)
            json_response = client.models.generate_content(
                model="gemini-2.0-flash-exp",
                contents=[json_prompt, myfile]
            )
            json_result = json_response.text
            json_response_path = save_response_to_file(json_result, "json_response")
            debug_output.append(f"JSON response received ({len(json_result)} characters)")
            extracted_json_data = None
            try:
                def extract_json_objects(text):
                    print(text)
                    match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
                    if not match:
                        raise ValueError("No JSON block found.")

                    json_block = match.group(1)
                    json_block = re.sub(r'\}\s*,?\s*\{', '},{', json_block)
                    json_block = f"[{json_block}]"

                    try:
                        parsed = json.loads(json_block)
                    except json.JSONDecodeError as e:
                        raise ValueError(f"Invalid JSON: {e}")

                    return parsed
                json_match = re.search(r'\[\s*\{.*\}\s*\]', json_result, re.DOTALL)
                if json_match:
                    json_text = json_match.group(0)
                else:
                    match = re.search(r'```(?:json)?\s*(.*?)```', json_result, re.DOTALL)
                    if match:
                        json_text = match.group(1)
                    else:
                        json_text = json_result

                raw_json_path = os.path.join(RESPONSE_LOGS_DIR, f"{timestamp}_raw_json.json")
                with open(raw_json_path, 'w', encoding='utf-8') as f:
                    f.write(json_text)
                response_log_files.append(raw_json_path)
                debug_output.append(f"Raw JSON saved to: {raw_json_path}")
            except Exception as e:
                debug_output.append(f"Error saving raw JSON: {str(e)}")
        
        except Exception as e:
            error_msg = f"Error generating JSON: {str(e)}"
            debug_output.append(error_msg)
            logger.error(error_msg)
            json_result = None
        def color_from_hex(hexstr):
            if isinstance(hexstr, list):
                return RGBColor(*hexstr)
            if not hexstr or hexstr.lower() == "auto":
                return None
            hexstr = hexstr.lstrip('#')
            return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))
        def json_to_docx(json_data, docx_path, template_path=None):
            if template_path:
                shutil.copyfile(template_path, docx_path)
                doc = Document(docx_path)
                for element in reversed(doc.paragraphs):
                    element.clear()
            else:
                doc = Document()
            for para in json_data:
                new_para = doc.add_paragraph(style=para.get("style", "Normal"))
                for run_info in para["runs"]:
                    run = new_para.add_run(run_info["text"])
                    run.bold = run_info.get("bold", False)
                    run.italic = run_info.get("italic", False)
                    if "color" in run_info:
                        run.font.color.rgb = color_from_hex(run_info["color"])
            print("YESYSYYSYS")

            doc.save(docx_path)
        data = extract_json_objects(json_result)
        for item in data:
            if "style" in item and isinstance(item["style"], str):
                item["style"] = item["style"].capitalize()  
        json_to_docx(
            data,
            output_path,
        )
        print('DONE')
        debug_output.append(f"Document saved with response information")
        
        # Render debug page or send file
        if debug_mode:
            debug_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Response Files</title>
                <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 20px; }}
                    pre {{ background-color: #f5f5f5; padding: 15px; border-radius: 5px; overflow-x: auto; }}
                    .card {{ margin-bottom: 20px; }}
                </style>
            </head>
            <body>
                <div class="container mt-5">
                    <h1>API Responses Saved</h1>
                    <div class="alert alert-success">
                        <p>All Gemini API responses have been saved to files for your review.</p>
                    </div>
                    
                    <div class="mb-3">
                        <a href="/download/{os.path.basename(output_path)}" class="btn btn-primary">Download Summary Document</a>
                        <a href="/" class="btn btn-secondary ms-2">Back to Form</a>
                    </div>
                    
                    <div class="card">
                        <div class="card-header">
                            <h3>Response Files</h3>
                        </div>
                        <div class="card-body">
                            <ul class="list-group">
                                {os.linesep.join([f'<li class="list-group-item"><a href="/view/{os.path.basename(path)}">{os.path.basename(path)}</a> - {path}</li>' for path in response_log_files])}
                            </ul>
                        </div>
                    </div>
                    
                    <div class="card">
                        <div class="card-header">
                            <h3>Processing Log</h3>
                        </div>
                        <div class="card-body">
                            <pre>{os.linesep.join(debug_output)}</pre>
                        </div>
                    </div>
                </div>
            </body>
            </html>
            """
            
            return render_template_string(debug_html)
        else:
            # Send file directly
            return send_file(output_path, as_attachment=True, download_name='responses_summary.docx')
            
    except Exception as e:
        error_message = f"Error: {str(e)}"
        logger.error(error_message)
        
        error_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Error</title>
            <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
        </head>
        <body>
            <div class="container mt-5">
                <div class="alert alert-danger">
                    <h4>Error Occurred</h4>
                    <p>{error_message}</p>
                </div>
                
                <div class="card mt-3">
                    <div class="card-header">Debug Information</div>
                    <div class="card-body">
                        <pre>{os.linesep.join(debug_output)}</pre>
                    </div>
                </div>
                
                <a href="/" class="btn btn-primary mt-3">Back to Form</a>
            </div>
        </body>
        </html>
        """
        
        return render_template_string(error_html), 500
    
    finally:
        # In a production environment, you would uncomment this to clean up
        # For debugging purposes, we're leaving the temp files
        # shutil.rmtree(temp_dir, ignore_errors=True)
        pass

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    # This route is used to download the generated file
    for root, _, files in os.walk(tempfile.gettempdir()):
        if filename in files:
            return send_file(os.path.join(root, filename), as_attachment=True)
    return "File not found", 404

@app.route('/view/<filename>', methods=['GET'])
def view_file(filename):
    # This route is used to view saved response files in the browser
    filepath = os.path.join(RESPONSE_LOGS_DIR, filename)
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Display the file content in a nice format
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>View File: {filename}</title>
            <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
            <style>
                pre {{ background-color: #f5f5f5; padding: 15px; border-radius: 5px; overflow-x: auto; }}
            </style>
        </head>
        <body>
            <div class="container mt-5">
                <h1>File: {filename}</h1>
                <div class="mb-3">
                    <a href="/download_response/{filename}" class="btn btn-primary">Download File</a>
                    <a href="javascript:history.back()" class="btn btn-secondary ms-2">Back</a>
                </div>
                <div class="card">
                    <div class="card-header">Content</div>
                    <div class="card-body">
                        <pre>{content}</pre>
                    </div>
                </div>
            </div>
        </body>
        </html>
        """
        return render_template_string(html)
    else:
        return "File not found", 404

@app.route('/download_response/<filename>', methods=['GET'])
def download_response_file(filename):
    # This route is used to download saved response files
    filepath = os.path.join(RESPONSE_LOGS_DIR, filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    return "File not found", 404

def secure_filename(filename):
    """Make the filename secure by removing path info"""
    return os.path.basename(filename)

if __name__ == '__main__':
    app.run(debug=True)