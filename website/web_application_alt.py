import os
import re
import json
import tempfile
import shutil
import logging
import time
from flask import Flask, request, render_template_string, send_file
from docx import Document
from docx.shared import RGBColor, Pt
import pathlib
import datetime

# Import Google GenAI
from google import genai
from google.genai import types

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# Create a permanent directory for response logs
RESPONSE_LOGS_DIR = "gemini_responses"
os.makedirs(RESPONSE_LOGS_DIR, exist_ok=True)

# HTML template for the form
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CACI Form Generator</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <style>
        body {
            background-color: #f5f5f5;
            font-family: Arial, sans-serif;
        }
        .container {
            max-width: 800px;
            margin: 50px auto;
            background-color: white;
            border-radius: 10px;
            padding: 30px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }
        .header {
            text-align: center;
            margin-bottom: 30px;
        }
        .upload-section {
            border: 2px dashed #ccc;
            border-radius: 8px;
            padding: 20px;
            text-align: center;
            margin-bottom: 20px;
        }
        .submit-btn {
            width: 100%;
            padding: 12px;
            font-size: 16px;
            margin-top: 20px;
        }
        pre {
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
            font-size: 12px;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>CACI Form Generator</h1>
            <p class="text-muted">Upload your legal documents to generate filled CACI forms</p>
        </div>

        <form action="/process" method="post" enctype="multipart/form-data">
            <div class="row">
                <div class="col-md-6">
                    <div class="upload-section">
                        <h4>LA CIV 244 PDF</h4>
                        <p>Upload the LA CIV 244 document</p>
                        <input type="file" class="form-control" name="laciv_file" accept=".pdf" required>
                    </div>
                </div>
                <div class="col-md-6">
                    <div class="upload-section">
                        <h4>Case Details PDF</h4>
                        <p>Upload the UD 105 document</p>
                        <input type="file" class="form-control" name="case_file" accept=".pdf" required>
                    </div>
                </div>
            </div>

            <div class="upload-section mt-3">
                <h4>CACI Template (DOCX)</h4>
                <p>Upload the Jury Instruction Template document</p>
                <input type="file" class="form-control" name="template_file" accept=".txt" required>
            </div>

            <div class="form-check mt-3">
                <input class="form-check-input" type="checkbox" id="save-responses" name="save_responses" checked>
                <label class="form-check-label" for="save-responses">
                    Save Gemini API responses to files
                </label>
            </div>
            
            <div class="form-check mt-2">
                <input class="form-check-input" type="checkbox" id="debug-mode" name="debug_mode" checked>
                <label class="form-check-label" for="debug-mode">
                    Show processing details
                </label>
            </div>

            <button type="submit" class="btn btn-primary submit-btn">Generate CACI Forms</button>
        </form>
    </div>
</body>
</html>
"""

@app.route('/', methods=['GET'])
def index():
    # Render the HTML form
    return render_template_string(HTML_TEMPLATE)

@app.route('/process', methods=['POST'])
def process():
    # Create a temporary directory for all processing
    temp_dir = tempfile.mkdtemp()
    debug_mode = 'debug_mode' in request.form
    save_responses = 'save_responses' in request.form
    debug_output = []
    
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    response_log_files = []

    
    def save_response_to_file(response_text, label):
        import json
        if save_responses:
            filename = f"{timestamp}_{label}.txt"
            filepath = os.path.join(RESPONSE_LOGS_DIR, filename)
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(response_text)
            response_log_files.append(filepath)
            debug_output.append(f"Saved {label} response to: {filepath}")
            return filepath
        return None
    
    try:
        # Save uploaded files to temp directory
        
        laciv_file = request.files['laciv_file']
        case_file = request.files['case_file']
        template_file = request.files['template_file']
        
        laciv_path = os.path.join(temp_dir, secure_filename(laciv_file.filename))
        case_path = os.path.join(temp_dir, secure_filename(case_file.filename))
        template_path = os.path.join(temp_dir, secure_filename(template_file.filename))
        
        laciv_file.save(laciv_path)
        case_file.save(case_path)
        template_file.save(template_path)
        
        debug_output.append(f"Files saved successfully")
        
        # Output paths
        output_path = os.path.join(temp_dir, 'final_output.docx')
        from docx import Document
        from docx.shared import RGBColor
        import json

        from google import genai
        from google.genai import types
        import pathlib
        import os
        import PIL
        # Set up the API key and client
        api_key = "AIzaSyAQtSELrheQ1szjk8oBBJAxq0RXgRdpWvg"
        client = genai.Client(api_key=api_key)
        filepath = pathlib.Path(case_path)
        prompt = 'Give me all information you have about the case'
        response = client.models.generate_content(
            model="gemini-2.0-flash-exp",
            contents=[
                types.Part.from_bytes(
                    data=filepath.read_bytes(),
                    mime_type='application/pdf',
                ),
                prompt
            ]
        )
        cont = response.text
        def generate_content():
            api_key = "AIzaSyAQtSELrheQ1szjk8oBBJAxq0RXgRdpWvg"
            client = genai.Client(api_key=api_key)
            myfile = client.files.upload(file=template_path)
            # Build prompt with formatting context
            prompt = f"""Legal Document Generation Task:
    Fill in the CACI template for this given CACI form using the provided context: "CACI Number": "101", "Jury Instructions": "Overview of Trial". Use this as context: {cont} to fill the form placeholders. Use the context and fill in placeholder such as defendant/plaintiff/city/address. Make your best guess for the placeholders. I want to see all placeholders filled. Dont output anything else. Keep the output in the same format as the input. JSON format, with the same attributes as the provided doc, you only need to change the text as you see fit.
    Use the same attributes as the provided text with the same styling. I need the same format as input. JSON format with the style, run, text, everything in style formatting. Fill the entire CACI form.  I need the entire form bro, i dont get what you are not understanding. I only need the CACI 101 form: Overview of a trial. I dont need superflous stuff
    Ignore any lines of this form. I only need the CACI 101:Overview of Trial. Your response should be one JSON file only with records being separated by commas. Only the form style json, no placeholders, nothing apart from the filled form w style details
    "style": "normal",
    "runs": [
      
        "text": "",
        "bold": null,
        "italic": null,
        "color": null
      
    ]
    Output:"""
            
            response = client.models.generate_content(
                    model="gemini-2.0-flash-exp",
                    contents=[
                        prompt, myfile
                    ]
            )
            
            return response.text


        import json
        from docx import Document
        from docx.shared import RGBColor
        import shutil
        import re

        def extract_json_objects(text):
            print(text)
            match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
            if not match:
                raise ValueError("No JSON block found.")

            json_block = match.group(1)
            json_block = re.sub(r'\}\s*,?\s*\{', '},{', json_block)
            json_block = f"[{json_block}]"

            try:
                parsed = json.loads(json_block)
            except json.JSONDecodeError as e:
                raise ValueError(f"Invalid JSON: {e}")

            return parsed

        def get_first_10_records(text):
            data = extract_json_objects(text)
            return data

        def color_from_hex(hexstr):
            if isinstance(hexstr, list):
                return RGBColor(*hexstr)
            if not hexstr or hexstr.lower() == "auto":
                return None
            hexstr = hexstr.lstrip('#')
            return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))

        def json_to_docx(json_data, docx_path, template_path=None):
            if template_path:
                shutil.copyfile(template_path, docx_path)
                doc = Document(docx_path)
                for element in reversed(doc.paragraphs):
                    element.clear()
            else:
                doc = Document()

            for para in json_data:
                new_para = doc.add_paragraph(style=para.get("style", "Normal"))
                for run_info in para["runs"]:
                    run = new_para.add_run(run_info["text"])
                    run.bold = run_info.get("bold", False)
                    run.italic = run_info.get("italic", False)
                    if "color" in run_info:
                        run.font.color.rgb = color_from_hex(run_info["color"])
            print("YESYSYYSYS")

            doc.save(docx_path)
        try:
            result = generate_content()
            with open("debug.txt", "w") as file:
                file.write(str(result))
            data = get_first_10_records(result)
            print(data)
            print("TYPE = ", type(data))

            for item in data:
                if "style" in item and isinstance(item["style"], str):
                    item["style"] = item["style"].capitalize()
            print(data)    
            json_to_docx(
                data,
                output_path,
            )
            print('DONE')
            return send_file(output_path, as_attachment=True, download_name='responses_summary.docx')

        except json.JSONDecodeError as e:
            print(f"Invalid JSON: {str(e)}")
            print("Raw LLM output:")
            print(result)
        except Exception as e:
            print(f"Error: {str(e)}")
    except:
        print("FAILEDDD")

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    # This route is used to download the generated file
    for root, _, files in os.walk(tempfile.gettempdir()):
        if filename in files:
            return send_file(os.path.join(root, filename), as_attachment=True)
    return "File not found", 404

@app.route('/view/<filename>', methods=['GET'])
def view_file(filename):
    # This route is used to view saved response files in the browser
    filepath = os.path.join(RESPONSE_LOGS_DIR, filename)
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Display the file content in a nice format
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>View File: {filename}</title>
            <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
            <style>
                pre {{ background-color: #f5f5f5; padding: 15px; border-radius: 5px; overflow-x: auto; }}
            </style>
        </head>
        <body>
            <div class="container mt-5">
                <h1>File: {filename}</h1>
                <div class="mb-3">
                    <a href="/download_response/{filename}" class="btn btn-primary">Download File</a>
                    <a href="javascript:history.back()" class="btn btn-secondary ms-2">Back</a>
                </div>
                <div class="card">
                    <div class="card-header">Content</div>
                    <div class="card-body">
                        <pre>{content}</pre>
                    </div>
                </div>
            </div>
        </body>
        </html>
        """
        return render_template_string(html)
    else:
        return "File not found", 404

@app.route('/download_response/<filename>', methods=['GET'])
def download_response_file(filename):
    # This route is used to download saved response files
    filepath = os.path.join(RESPONSE_LOGS_DIR, filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    return "File not found", 404

def secure_filename(filename):
    """Make the filename secure by removing path info"""
    return os.path.basename(filename)

if __name__ == '__main__':
    app.run(debug=True)